import PyPDF2
import docx
import os
from typing import Optional

class TextExtractor:
    """Extract text from different file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text or None if error occurs
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
            
            return text.strip()
        
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            return None
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text or None if error occurs
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            return None
    
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """
        Extract text from file based on extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or None if error occurs
        """
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return TextExtractor.extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return TextExtractor.extract_from_docx(file_path)
        else:
            print(f"Unsupported file format: {file_extension}")
            return None